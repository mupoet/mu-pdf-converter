#!/usr/bin/env python3
"""
pdf_to_docx.py - 将 PDF 转换为可编辑的 Word 文档
用法：python pdf_to_docx.py input.pdf [--outfile output.docx]

转换策略：
- 用 PyMuPDF 提取文本块（含字体大小、加粗、斜体）
- 用 pdfplumber 识别表格区域（避免重复提取）
- 标题判断：字号比正文大 2pt 以上 → Heading 1
- 表格用 python-docx Table 对象插入，首行加粗
"""

import argparse
import sys
from pathlib import Path

# ── 依赖检查 ──────────────────────────────────
_missing = []
try:
    from markitdown.converters._pdf_converter import _extract_form_content_from_words as _markitdown_extract
    _HAS_MARKITDOWN = True
except ImportError:
    _HAS_MARKITDOWN = False

try:
    import fitz  # PyMuPDF
except ImportError:
    _missing.append("pymupdf")
try:
    import pdfplumber
except ImportError:
    _missing.append("pdfplumber")
try:
    from docx import Document
    from docx.shared import Pt, RGBColor as DocxRGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import docx.oxml as oxml
except ImportError:
    _missing.append("python-docx")

if _missing:
    print(
        f"[错误] 缺少依赖：{', '.join(_missing)}\n"
        f"请运行：pip install {' '.join(_missing)}",
        file=sys.stderr,
    )
    sys.exit(1)

sys.path.insert(0, str(Path(__file__).parent))
from utils import (  # noqa: E402
    generate_output_path,
    is_in_table_region,
    fitz_color_to_rgb,
    map_font_name,
    is_scanned_pdf,
    check_scanned_and_warn,
    parse_markdown_tables,
)


# _parse_markdown_tables 已移至 utils.py 公共模块（parse_markdown_tables）


# ─────────────────────────────────────────────
# 核心逻辑
# ─────────────────────────────────────────────

def pdf_to_docx(
    pdf_path: str,
    outfile: str = None,
) -> str:
    """
    将 PDF 转换为 Word 文档。

    参数:
        pdf_path: 输入 PDF 路径
        outfile: 输出 docx 路径（可选）
    返回:
        输出文件路径
    """
    pdf_path = str(Path(pdf_path).resolve())
    if not Path(pdf_path).exists():
        print(f"[错误] 文件不存在：{pdf_path}", file=sys.stderr)
        sys.exit(1)

    output_path = generate_output_path(pdf_path, ".docx", outfile=outfile)

    # 检测是否为扫描件
    if is_scanned_pdf(pdf_path):
        print("[警告] 检测到扫描件（无文本层），转换结果可能为空", file=sys.stderr)
        print("       建议：先用 pdf_to_images.py 转为图片，再使用 OCR 工具", file=sys.stderr)

    doc = Document()
    fitz_doc = fitz.open(pdf_path)
    total_pages = len(fitz_doc)
    print(f"[信息] 共 {total_pages} 页，开始转换为 Word...")

    # 统计全文基准字体大小（用于标题判断）
    base_font_size = _detect_base_font_size(fitz_doc)
    print(f"[信息] 检测到基准字体大小：{base_font_size:.1f}pt")

    with pdfplumber.open(pdf_path) as plumber_pdf:
        for page_idx in range(total_pages):
            fitz_page = fitz_doc[page_idx]
            plumber_page = plumber_pdf.pages[page_idx]
            page_num = page_idx + 1

            # 获取表格区域 bbox 与表格数据：优先用 MarkItDown 识别无边框表格
            markitdown_tables = None
            if _HAS_MARKITDOWN:
                try:
                    md_result = _markitdown_extract(plumber_page)
                    if md_result is not None:
                        markitdown_tables = parse_markdown_tables(md_result)
                        if not markitdown_tables:
                            markitdown_tables = None
                except Exception:
                    markitdown_tables = None

            if markitdown_tables is not None:
                # 无边框表格路径：表格区域 bbox 无法精确定位，设为空（可接受）
                table_bboxes = []
                tables_data = markitdown_tables
            else:
                # 降级：原有 pdfplumber 有边框表格路径
                table_bboxes = _get_table_bboxes(plumber_page)
                tables_data = plumber_page.extract_tables()

            # 提取文本块（按阅读顺序排序）
            text_dict = fitz_page.get_text("dict", sort=True)
            blocks = text_dict.get("blocks", [])

            # 插入页分隔符（第一页除外）
            if page_idx > 0:
                doc.add_page_break()

            # 处理每个文本块
            table_inserted_bboxes = set()
            for block in blocks:
                if block.get("type") != 0:  # 0=文本，1=图片
                    continue

                block_bbox = block.get("bbox", (0, 0, 0, 0))

                # 跳过表格区域的文本块（由表格处理逻辑统一插入）
                if is_in_table_region(block_bbox, table_bboxes):
                    continue

                # 合并块内所有 span 的文本
                lines_text = []
                for line in block.get("lines", []):
                    line_parts = []
                    for span in line.get("spans", []):
                        line_parts.append(span.get("text", ""))
                    lines_text.append("".join(line_parts))
                full_text = "\n".join(lines_text).strip()

                if not full_text:
                    continue

                # 获取块的主要字体属性（取第一个 span）
                first_span = _get_first_span(block)
                font_size = first_span.get("size", 11.0) if first_span else 11.0
                is_bold = bool(first_span.get("flags", 0) & 0b10000) if first_span else False
                is_italic = bool(first_span.get("flags", 0) & 0b01) if first_span else False
                color = fitz_color_to_rgb(first_span.get("color", 0)) if first_span else (0, 0, 0)
                font_name = map_font_name(first_span.get("font", "")) if first_span else "Arial"

                # 判断是否为标题
                is_heading = font_size >= base_font_size + 2.0

                if is_heading:
                    para = doc.add_heading(full_text, level=1)
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(full_text)
                    run.bold = is_bold
                    run.italic = is_italic
                    run.font.size = Pt(font_size)
                    run.font.name = font_name
                    try:
                        run.font.color.rgb = DocxRGBColor(*color)
                    except Exception:
                        pass

            # 插入表格（按页面顺序，在文本块后）
            for t_idx, table_data in enumerate(tables_data):
                if not table_data:
                    continue
                _insert_table(doc, table_data)

            plumber_page.close()  # 修复内存泄漏
            print(f"  ✓ 第 {page_num} 页处理完成（表格 {len(tables_data)} 个）")

    fitz_doc.close()
    doc.save(output_path)
    print(f"\n[完成] Word 文档已保存 → {output_path}")
    return output_path


def _detect_base_font_size(fitz_doc, sample_pages: int = 5) -> float:
    """
    检测文档基准字体大小（最频繁出现的字号）。
    用于区分正文和标题。
    """
    from collections import Counter
    size_counter = Counter()
    pages_to_check = min(sample_pages, len(fitz_doc))

    for i in range(pages_to_check):
        page = fitz_doc[i]
        text_dict = page.get_text("dict")
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = round(span.get("size", 11.0))
                    text = span.get("text", "").strip()
                    if text and size > 0:
                        size_counter[size] += len(text)

    if not size_counter:
        return 11.0
    return float(size_counter.most_common(1)[0][0])


def _get_first_span(block: dict) -> dict:
    """从 fitz 文本块中获取第一个 span"""
    for line in block.get("lines", []):
        for span in line.get("spans", []):
            return span
    return {}


def _get_table_bboxes(plumber_page) -> list:
    """从 pdfplumber 页面获取所有表格的 bbox 列表"""
    bboxes = []
    try:
        for table in plumber_page.find_tables():
            bbox = table.bbox  # (x0, top, x1, bottom)
            bboxes.append(bbox)
    except Exception:
        pass
    return bboxes


def _insert_table(doc: "Document", table_data: list):
    """在 Word 文档中插入表格"""
    if not table_data:
        return

    rows = len(table_data)
    # 计算最大列数（防止不规则表格）
    cols = max((len(row) for row in table_data), default=1)
    if cols == 0:
        return

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(table_data):
        row = table.rows[r_idx]
        for c_idx, cell_val in enumerate(row_data):
            if c_idx >= cols:
                break
            cell_text = str(cell_val or "").replace("\n", " ").strip()
            cell = row.cells[c_idx]
            cell.text = cell_text
            # 首行加粗
            if r_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                    if not para.runs and cell_text:
                        run = para.add_run(cell_text)
                        run.bold = True
                        # 清除之前直接赋值的文本
                        cell.paragraphs[0].clear()
                        cell.paragraphs[0].add_run(cell_text).bold = True
                        break

    doc.add_paragraph()  # 表格后加空行


# ─────────────────────────────────────────────
# CLI 入口
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="将 PDF 转换为可编辑的 Word 文档（.docx）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python pdf_to_docx.py report.pdf
  python pdf_to_docx.py report.pdf --outfile output/result.docx

注意:
  - 扫描件 PDF（无文本层）转换结果可能为空，建议先 OCR 处理
  - 复杂排版（多栏、浮动图片等）可能无法完美还原
  - 表格识别依赖 pdfplumber，有边框的表格识别率更高
        """,
    )
    parser.add_argument("input", help="输入 PDF 文件路径")
    parser.add_argument(
        "--outfile", default=None,
        help="输出 docx 文件路径（默认与 PDF 同名同目录）"
    )

    args = parser.parse_args()
    check_scanned_and_warn(args.input)
    pdf_to_docx(
        pdf_path=args.input,
        outfile=args.outfile,
    )


if __name__ == "__main__":
    main()
